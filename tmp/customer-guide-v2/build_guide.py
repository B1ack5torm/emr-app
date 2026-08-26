from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Admin\Documents\emr-app")
IMG = ROOT / "tmp" / "customer-guide-v2" / "sanitized"
OUT = ROOT / "output" / "docx" / "CareChart_Customer_Configuration_and_User_Guide.docx"

# compact_reference_guide with one named override: CareChart brand green.
INK = "17252B"
GREEN = "2F6F5E"
DARK_GREEN = "194B3E"
MUTED = "5D716B"
PALE_GREEN = "E8F2EE"
CREAM = "F8F5ED"
PALE_GOLD = "FFF4DD"
PALE_RED = "FBE9E4"
BLUE_GRAY = "E8EEF5"
BORDER = "D9DED8"
WHITE = "FFFFFF"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])


def set_list_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def restart_numbering(doc, base_num_id=5):
    numbering = doc.part.numbering_part.element
    base = next(item for item in numbering.findall(qn("w:num")) if int(item.get(qn("w:numId"))) == base_num_id)
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_id = max([int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))] or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, DARK_GREEN, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("CareChart Customer Guide")
    set_font(r, 9, GREEN, True)
    r = hp.add_run("  |  Configuration and daily use")
    set_font(r, 9, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Page ")
    set_font(r, 9, MUTED)
    add_page_field(fp)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CUSTOMER ENABLEMENT GUIDE")
    set_font(r, 10, GREEN, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CareChart Configuration and User Guide")
    set_font(r, 30, DARK_GREEN, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("A plain-language handbook for hospital setup, staff training, and daily operations")
    set_font(r, 14, MUTED)

    table = doc.add_table(rows=2, cols=2)
    data = [
        ("Prepared for", "Hospital administrators and authorized staff"),
        ("Edition", "Customer edition - August 2026"),
        ("Scope", "CareChart outpatient workflows"),
        ("Screenshots", "Demonstration data only; identifiers removed or replaced"),
    ]
    for idx, (label, value) in enumerate(data):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        cell_shading(cell, CREAM)
        p = cell.paragraphs[0]
        rr = p.add_run(label.upper() + "\n")
        set_font(rr, 8, GREEN, True)
        rr = p.add_run(value)
        set_font(rr, 10, INK, True)
    set_table_geometry(table, [4680, 4680], indent=120)

    doc.add_paragraph()
    add_callout(
        doc,
        "Privacy-safe customer edition",
        "The screenshots in this guide are flattened demonstration images. Patient names, MRNs, phone numbers, dates of birth, staff emails, order identifiers, clinical narratives, payment links, and audit identifiers were removed or replaced before the images were embedded.",
        PALE_GREEN,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CareChart · Care, connected")
    set_font(r, 12, GREEN, True)
    doc.add_page_break()


def add_callout(doc, label_text, body, fill=BLUE_GRAY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label_text + ": ")
    set_font(r, 10.5, DARK_GREEN, True)
    r = p.add_run(body)
    set_font(r, 10.5, INK)
    set_table_geometry(table, [9360], indent=120)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_list_numbering(p, 1)
        p.add_run(item)


def add_steps(doc, items):
    num_id = restart_numbering(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_list_numbering(p, num_id)
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_shading(cell, PALE_GREEN)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, 10, DARK_GREEN, True)
    repeat_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, 9.5, INK)
    set_table_geometry(table, widths, indent=120)
    return table


def set_alt(inline_shape, text):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", text)


def add_screen(doc, name, caption, alt, width=6.35):
    path = IMG / f"{name}.png"
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_alt(shape, alt)
    cp = doc.add_paragraph(caption, style="Caption")
    return cp


def add_screen_pair(doc, left, right, caption, left_alt, right_alt, left_width=3.05, right_width=3.05):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    l = p.add_run().add_picture(str(IMG / f"{left}.png"), width=Inches(left_width))
    set_alt(l, left_alt)
    spacer = p.add_run("   ")
    set_font(spacer, 8)
    r = p.add_run().add_picture(str(IMG / f"{right}.png"), width=Inches(right_width))
    set_alt(r, right_alt)
    doc.add_paragraph(caption, style="Caption")


def h1(doc, text, new_page=False):
    p = doc.add_paragraph(text, style="Heading 1")
    if new_page:
        p.paragraph_format.page_break_before = True
    return p


def h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def add_toc_placeholder(doc):
    h1(doc, "Contents")
    p = doc.add_paragraph("[[TOC]]")
    p.paragraph_format.space_after = Pt(10)
    add_callout(doc, "How to use this guide", "Administrators should begin with hospital setup and role access. Other staff can go directly to the section matching their daily work.", CREAM)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_toc_placeholder(doc)

    h1(doc, "1. What CareChart does")
    doc.add_paragraph(
        "CareChart brings common outpatient hospital tasks into one hospital-scoped workspace. It connects patient registration, appointments, the doctor's queue, clinical documentation, diagnostic orders, patient records, documents, billing, and audit history."
    )
    add_screen(doc, "01_landing", "Figure 1. Public landing page with links for booking, patient access, and staff sign-in.", "CareChart public landing page with navigation for features, patient portal, appointment booking, and sign in.", 6.2)
    h2(doc, "The top navigation")
    add_table(doc, ["Area", "Plain-language purpose"], [
        ("Front Desk", "Find or register patients and start walk-in visits."),
        ("Schedule", "Book appointments and check patients in."),
        ("Doctor's Desk", "Review requests, open the waiting room, and document visits."),
        ("Patient Records", "Search the hospital's patient history."),
        ("Admin", "Create hospital accounts, invite staff, and manage roles."),
        ("Billing", "Generate invoices and record payments, refunds, and voids."),
        ("Diagnostics", "Start lab/radiology work, upload reports, and track review."),
        ("Documents", "Upload approved patient PDFs and images."),
        ("Audit", "Review important security and workflow events."),
        ("Settings", "Configure clinics, practitioners, schedules, prices, and tax."),
    ], [2000, 7360])

    h1(doc, "2. Sign in, passwords, and the active hospital")
    h2(doc, "Sign in")
    add_steps(doc, [
        "Open the staff sign-in page.",
        "Enter the email address assigned to your individual account.",
        "Enter your password and select Sign in.",
        "If you cannot remember the password, use Forgot password? and follow the approved reset process.",
    ])
    add_screen(doc, "02_login", "Figure 2. Staff sign-in screen using demonstration credentials.", "CareChart staff sign-in form with email, password, forgot-password link, and sign-in button.", 3.25)
    add_callout(doc, "Password rule", "A new password must contain at least 8 characters, including at least one letter and one number. Never share a password or use one account for several people.", PALE_GOLD)
    h2(doc, "Active hospital")
    doc.add_paragraph("The hospital shown at the top of the screen is the active hospital. Data entered or viewed on that screen belongs to that hospital workspace.")
    add_bullets(doc, [
        "Super Admin users can open Active Hospital and switch between hospital accounts they manage.",
        "Other staff see only the hospital or hospitals to which they are assigned.",
        "Always confirm the hospital name before registering a patient, changing settings, or creating an invoice.",
    ])
    add_callout(doc, "Important", "Changing the active hospital changes the workspace. It does not move patients, visits, staff, invoices, or configuration between hospitals.", PALE_RED)

    h1(doc, "3. Roles and access")
    doc.add_paragraph("CareChart shows only the menus allowed for the signed-in role. An administrator should give each person the smallest role needed for their job.")
    add_table(doc, ["Role", "Typical access"], [
        ("Super Admin", "All hospital workspaces, hospital creation, active-hospital switching, and all operational areas."),
        ("Admin / Clinic Admin", "Hospital administration, settings, records, billing, diagnostics, documents, and audit."),
        ("Doctor", "Schedule, Doctor's Desk, records, diagnostics, and documents."),
        ("Front Desk / Reception", "Patient registration, schedule, records, billing support, and documents."),
        ("Nurse", "Schedule, patient records, and documents."),
        ("Billing", "Billing and patient-record lookup needed for finance work."),
        ("Lab / Radiology", "Diagnostics, patient-record lookup, and documents needed for report workflows."),
        ("Patient", "The separate patient portal connected only to that patient's account."),
    ], [2200, 7160])
    add_callout(doc, "Role changes", "Role changes take effect for future access. Remove or suspend access promptly when a person leaves or no longer needs the system.", PALE_GREEN)

    h1(doc, "4. Front Desk: find or register a patient", new_page=True)
    h2(doc, "Find an existing patient")
    add_steps(doc, [
        "Open Front Desk.",
        "Search using the MRN, patient name, or phone number.",
        "Check more than one detail before selecting a result. This helps prevent choosing the wrong person.",
        "If no correct record exists, select Register new patient.",
    ])
    add_screen(doc, "03_front_desk", "Figure 3. Front Desk search before a patient is selected.", "Front Desk page with patient search field and Register new patient button.")
    h2(doc, "Register a new patient")
    add_steps(doc, [
        "Enter the patient's full name exactly as provided.",
        "Enter date of birth when available; confirm that the calculated age is correct.",
        "Choose gender and the doctor responsible for the visit.",
        "Enter phone, email, blood group, emergency contact, and address when collected under hospital policy.",
        "Add known allergies carefully. Press Enter after each allergy so it is saved as a separate item.",
        "Review the information with the patient, then select Save patient & continue to visit.",
    ])
    add_screen(doc, "04_registration", "Figure 4. Empty New Patient Registration form.", "New patient registration form showing identity, contact, doctor, blood group, address, and allergy fields.", 4.2)
    add_callout(doc, "Duplicate prevention", "Search before registering. If two records may belong to the same person, stop and ask an administrator to use the approved duplicate-review process. Do not casually delete a patient record.", PALE_GOLD)

    h1(doc, "5. Appointments and check-in", new_page=True)
    h2(doc, "Book from the schedule")
    add_steps(doc, [
        "Open Schedule and choose the required date.",
        "Select New appointment.",
        "Search for and select the patient.",
        "Choose the doctor, time, visit duration, and reason.",
        "Select Book appointment.",
        "When the patient arrives, open the appointment and complete check-in so the visit reaches the doctor's queue.",
    ])
    add_screen(doc, "05_schedule", "Figure 5. Appointment Schedule showing an empty day.", "Appointment schedule with date controls and New appointment button.")
    add_screen(doc, "06_new_appointment", "Figure 6. New appointment form.", "New appointment form with patient search, doctor, time, duration, and reason fields.", 3.6)
    h2(doc, "Online appointment requests")
    add_bullets(doc, [
        "A public request is not a completed visit. It must be reviewed and confirmed by authorized staff.",
        "The doctor or authorized team reviews the requested date and time before confirmation.",
        "After confirmation, Front Desk checks the patient in when they arrive.",
    ])

    h1(doc, "6. Doctor's Desk and waiting room", new_page=True)
    doc.add_paragraph("Doctor's Desk brings together appointment requests, confirmed online appointments, and patients waiting for consultation.")
    add_steps(doc, [
        "Review any appointment requests awaiting action.",
        "Confirm or decline according to the hospital's booking process.",
        "Open a patient from Waiting Room only when ready to begin the consultation.",
        "Confirm the patient and visit reason before documenting anything.",
    ])
    add_screen(doc, "07_doctors_desk", "Figure 7. Doctor's Desk with a demonstration waiting-room entry.", "Doctor's Desk page showing appointment request areas and a demonstration patient in the waiting room.")

    h1(doc, "7. Document a clinical visit", new_page=True)
    h2(doc, "Complete the consultation")
    add_steps(doc, [
        "Confirm the patient header and review any allergy alert before prescribing.",
        "Read the chief complaint and record measured vital signs.",
        "Enter examination findings or clinical feedback.",
        "Record the diagnosis and advice given to the patient.",
        "Add medicines and tests when clinically appropriate.",
        "Create a tracked diagnostic order when the lab or radiology team must receive and process the request.",
        "Save a draft if the visit is not ready. Otherwise confirm the signature checkbox and select Complete & sign visit.",
    ])
    add_screen(doc, "08_consultation", "Figure 8. Consultation form using demonstration text only.", "Consultation form with demonstration patient header, complaint, vitals, examination, diagnosis, advice, and prescription section.")
    add_callout(doc, "Clinical responsibility", "CareChart records information and displays alerts. It does not replace professional judgement. The signing clinician remains responsible for accuracy, appropriateness, and follow-up.", PALE_RED)
    h2(doc, "Tests, diagnostic orders, and signature")
    add_screen_pair(
        doc,
        "09_consultation_orders",
        "10_signature",
        "Figure 9. Left: tracked operational diagnostic orders. Right: draft and digital-signature controls.",
        "Consultation section for tests and operational diagnostic orders using demonstration order IDs.",
        "Consultation signature area with Save draft and Complete and sign visit buttons.",
        3.08,
        3.08,
    )
    add_bullets(doc, [
        "Tests ordered is a simple clinical list recorded in the visit.",
        "Operational diagnostic orders create work for the Diagnostics team and support report upload and review tracking.",
        "Save draft keeps the visit open. Complete & sign visit finalizes the encounter for downstream records and billing.",
        "A correction after signing should follow the approved amendment process so the original remains traceable.",
    ])

    h1(doc, "8. Patient Records", new_page=True)
    add_steps(doc, [
        "Open Patient Records.",
        "Search by MRN, patient name, or phone number.",
        "Select the correct patient after verifying more than one identifier.",
        "Review permitted demographics, allergies, previous visits, diagnoses, medicines, tests, documents, and signed information.",
    ])
    add_screen(doc, "11_patient_records", "Figure 10. Patient Records list using demonstration identities.", "Patient Records page showing four fictional demonstration patients and fictional MRNs.")
    add_callout(doc, "Minimum necessary access", "Open a patient record only when it is needed for your assigned work. Searching or viewing a record can be recorded in the audit history.", PALE_GOLD)

    h1(doc, "9. Hospital and staff administration", new_page=True)
    h2(doc, "Create a hospital account - Super Admin")
    add_steps(doc, [
        "Open Admin and confirm the active hospital shown in the header.",
        "Under Create Hospital Account, enter the hospital or clinic name.",
        "Enter the first administrator's name and email.",
        "Set a temporary password that follows the password rule.",
        "Select Create hospital.",
        "Open Active Hospital and switch into the new hospital before configuring it.",
    ])
    add_screen(doc, "12_admin", "Figure 11. Admin page for hospital creation, staff invitations, and user management.", "Admin page with empty hospital creation fields, invitation form, Add user control, and pending request area.")
    h2(doc, "Invite or add staff")
    add_bullets(doc, [
        "Invite by email when the person should create their own password through the approved invitation process.",
        "Add user manually only when the hospital's onboarding process requires an administrator-created account.",
        "Choose the role carefully. The role controls which menus and records the person can access.",
        "Administrator email and password fields should be entered manually; they must not copy the signed-in user's credentials automatically.",
    ])
    h2(doc, "Review active staff")
    add_screen(doc, "13_staff", "Figure 12. Active staff list using demonstration names and emails.", "Active staff list with fictional names, example.com email addresses, role selectors, and Remove buttons.", 6.1)
    add_steps(doc, [
        "Review active staff regularly.",
        "Change a role only after confirming the person's current duties.",
        "Use Remove when access should end, following the hospital's offboarding policy.",
        "Never reuse a former employee's account for a new employee.",
    ])

    h1(doc, "10. Billing, invoices, and payments")
    h2(doc, "Generate and review invoices")
    add_steps(doc, [
        "Complete and sign the clinical visit.",
        "Open Billing and choose the completed visit or invoice.",
        "Review the patient, invoice status, date, line items, unit price, discount, tax, and total.",
        "If a price is zero unexpectedly, stop before collecting money and configure the matching service price under Settings.",
        "Save changes only after the invoice is correct.",
    ])
    add_screen(doc, "14_billing", "Figure 13. Billing list using demonstration patients and dates.", "Billing page showing fictional invoices with paid, unpaid, and void statuses.")
    add_screen(doc, "15_invoice", "Figure 14. Demonstration invoice with one consultation line.", "Invoice screen using a fictional patient and invoice number, with quantity, unit price, discount, tax, and total.", 6.0)
    h2(doc, "Payments, refunds, print, and void")
    add_screen(doc, "16_payment", "Figure 15. Demonstration payment and refund section.", "Payments section with demonstration payment date, cash method, paid and remaining amounts, and refund fields.", 6.0)
    add_bullets(doc, [
        "Record payment: enter the amount, method, and optional reference used by the hospital.",
        "Partial payment: the invoice remains partly unpaid until the balance is collected.",
        "Refund: enter the refund amount and a clear reason; the original payment remains traceable.",
        "Void invoice: use only when the invoice should be cancelled. A reason is required.",
        "Print: produce a customer-facing invoice copy through the approved printing process.",
    ])

    h1(doc, "11. Diagnostics: collection, report upload, and doctor review", new_page=True)
    h2(doc, "Lab or radiology team")
    add_steps(doc, [
        "Open Diagnostics and filter by type or status when needed.",
        "For a blood test, begin the collection process when the sample is actually received. For imaging, begin processing according to the department workflow.",
        "Complete the test or imaging work in the hospital's approved laboratory or radiology system.",
        "When the final report is ready, attach the approved PDF, JPEG, or PNG file to the in-progress order.",
        "Select Upload report & send to doctor. The order becomes ready for clinical review.",
    ])
    add_screen(doc, "17_diagnostics", "Figure 16. Diagnostics page using fictional orders, patient, clinician, and report name.", "Diagnostics page with one reviewed demonstration order and one in-progress order ready for report upload.")
    h2(doc, "Ordering clinician")
    add_steps(doc, [
        "Open the completed order.",
        "Open and read the attached report.",
        "When clinically satisfied, acknowledge the report using the available review control.",
        "Arrange any required patient follow-up outside the status change itself.",
    ])
    add_callout(doc, "Status meaning", "Completed means a report was uploaded. Reviewed means an authorized clinician opened and acknowledged it. A status is not a substitute for reading the report.", PALE_GOLD)

    h1(doc, "12. Patient documents", new_page=True)
    add_steps(doc, [
        "Open Documents.",
        "Search for the patient and verify the correct record.",
        "Choose an approved PDF, JPEG, or PNG file.",
        "Link it to a diagnostic order when the document is the final report for that order; otherwise keep it as a general patient document.",
        "Select Upload securely.",
        "Authorized staff can later open or download the document within the same hospital workspace.",
    ])
    add_screen(doc, "18_documents", "Figure 17. Patient Documents search screen.", "Patient Documents page with a search field and privacy-focused description.")
    add_callout(doc, "Before uploading", "Confirm the patient, document type, and hospital. Do not upload unrelated files, passwords, exports containing multiple patients, or documents not approved by hospital policy.", PALE_RED)

    h1(doc, "13. Audit history", new_page=True)
    doc.add_paragraph("Audit history is an append-only record of important security and workflow events for the active hospital. It can show the time, action, resource, user or system identifier, and reason when one was supplied.")
    add_screen(doc, "19_audit", "Figure 18. Audit history using demonstration timestamps and identifiers.", "Audit history table with fictional timestamps, resources, users, and example actions.")
    add_bullets(doc, [
        "Use audit history for authorized operational, privacy, security, and support reviews.",
        "Do not put passwords, tokens, or unnecessary clinical details into an audit reason.",
        "An audit event shows that an action occurred; it does not by itself explain whether the action was clinically appropriate.",
    ])

    h1(doc, "14. Configure a hospital in the recommended order", new_page=True)
    add_callout(doc, "Start here", "Switch to the correct active hospital before every configuration step. Configuration belongs only to that hospital workspace.", PALE_GREEN)
    h2(doc, "Step 1 - Locations and departments")
    add_steps(doc, [
        "Open Settings.",
        "Enter the clinic name, a unique clinic code, address, and IANA timezone such as Asia/Kolkata.",
        "Select Add clinic.",
        "Choose the clinic, select Department, enter the department name, and add the configuration.",
        "Repeat for each real care location and department.",
    ])
    add_screen(doc, "20_clinic_settings", "Figure 19. Clinic, location, and department configuration.", "Settings page with fields for clinic name, code, address, timezone, clinic selection, configuration type, and name.")
    h2(doc, "Step 2 - Practitioners and weekly schedules")
    add_steps(doc, [
        "Select the doctor account and the clinic where the doctor works.",
        "Choose the department and enter specialty, qualification, and registration number when required.",
        "Select Save practitioner.",
        "Choose the practitioner and weekday, then enter start time, end time, and slot length.",
        "Add a break only when both break start and break end are entered.",
        "Select Save Weekly Schedule. Saving the same practitioner and weekday updates the existing entry instead of creating a duplicate.",
    ])
    add_screen(doc, "21_practitioner_settings", "Figure 20. Practitioner profile and weekly schedule configuration.", "Settings form for doctor, clinic, department, specialty, qualification, registration number, weekday, work times, slot duration, and break.", 6.0)
    h2(doc, "Step 3 - Blocked time and holidays")
    add_bullets(doc, [
        "Blocked practitioner time is for leave, meetings, or another period when one practitioner is unavailable.",
        "Clinic holiday is for a date when the selected clinic is closed.",
        "Enter a clear reason or holiday name so staff understand why slots are unavailable.",
    ])
    h2(doc, "Step 4 - Service catalog, prices, and tax")
    add_steps(doc, [
        "Choose the clinic.",
        "Enter a stable service code and service name that match how the service will appear on invoices.",
        "Choose the category, such as Consultation, Medicine, Test, or Imaging.",
        "Enter the clinic price. Mark Taxable service only when approved by the hospital's finance process.",
        "Select Save service and price.",
        "If tax applies, choose the clinic, enter the tax name and rate, set the effective date, and add the tax rule.",
    ])
    add_screen(doc, "22_service_settings", "Figure 21. Blocked time, holidays, service pricing, and tax configuration.", "Settings page for practitioner blocked time, clinic holidays, service catalog, clinic price, taxable service, and tax rule.", 5.9)
    add_callout(doc, "Why prices matter", "Invoice prices are populated by matching visit items to the active clinic's service catalog. A missing or mismatched service can leave the invoice price at zero.", PALE_GOLD)

    h1(doc, "15. Public booking and patient portal")
    h2(doc, "Public appointment booking")
    add_bullets(doc, [
        "The public page can link a patient to the appointment request form.",
        "The patient chooses an available hospital, clinic, doctor, date, and time according to the deployed booking workflow.",
        "A request enters the review queue and becomes confirmed only after authorized review.",
    ])
    h2(doc, "Patient portal")
    add_bullets(doc, [
        "The patient uses the separate patient portal entry point, not the staff sign-in form.",
        "Activation must follow the hospital's identity-verification process.",
        "A portal user should see only the patient record connected to that portal account.",
        "Staff must verify the patient's email and identity before sending or renewing access.",
    ])

    h1(doc, "16. New hospital setup checklist")
    add_bullets(doc, [
        "Hospital account created and selected as the active hospital",
        "Clinic locations, codes, addresses, and timezones entered",
        "Departments and appointment types created",
        "Doctor accounts linked to practitioner profiles",
        "Every working weekday, slot length, and optional break checked",
        "Planned leave, blocked periods, and clinic holidays entered",
        "Common consultations, medicines, tests, and imaging services priced for each clinic",
        "Tax rules reviewed and approved by the hospital's finance owner",
        "Staff invited with the smallest role required",
        "Email delivery, document storage, encryption, malware scanning, backups, and HTTPS tested",
        "A complete fictional-patient workflow tested before entering real patient information",
    ])
    add_callout(doc, "Go-live approval", "Technical availability does not replace hospital approval. Clinical, operational, privacy, legal, security, and finance owners should approve the configured workflow before real use.", PALE_RED)

    h1(doc, "17. Daily workflow checklists", new_page=True)
    h2(doc, "Walk-in visit")
    add_steps(doc, [
        "Front Desk searches for or registers the patient.",
        "Front Desk logs the visit and selects the doctor.",
        "Doctor opens the patient from Waiting Room.",
        "Doctor documents, orders, and signs the visit.",
        "Lab or radiology processes any tracked order and uploads the final report.",
        "Ordering clinician reviews and acknowledges the report.",
        "Billing reviews prices, generates the invoice, and records payment.",
    ])
    h2(doc, "Online appointment")
    add_steps(doc, [
        "Patient submits a public appointment request.",
        "Authorized staff or doctor reviews and confirms it.",
        "Front Desk checks the patient in on arrival.",
        "Clinical, diagnostic, record, and billing steps follow the same controlled workflow as a walk-in visit.",
    ])

    h1(doc, "18. Privacy, security, and safe documentation")
    add_bullets(doc, [
        "Use an individual account. Never share passwords or leave a signed-in device unattended.",
        "Confirm the active hospital and patient before viewing, entering, uploading, printing, or billing information.",
        "Access only the minimum information needed for assigned work.",
        "Do not include patient information in training material, sales documents, screenshots, support messages, or presentations unless formally approved and securely de-identified.",
        "Use fictional identities for demonstrations. A common-looking name is not automatically de-identified.",
        "Remove staff emails, internal IDs, tokens, and environment details from customer-facing screenshots.",
        "Use approved storage, encryption, malware scanning, backups, HTTPS, access review, incident response, and retention procedures.",
        "Follow applicable hospital policy and legal/regulatory requirements for the country and care setting.",
    ])
    add_callout(doc, "Screenshot rule", "A customer guide should contain only screenshots that were deliberately prepared for demonstration. Cropping is not enough when the original file still contains hidden or out-of-frame sensitive pixels; use a flattened sanitized copy.", PALE_RED)

    h1(doc, "19. Troubleshooting", new_page=True)
    add_table(doc, ["Problem", "What to check"], [
        ("Wrong hospital is shown", "Super Admin: open Active Hospital and switch. Other staff: confirm assignment to the correct hospital."),
        ("A menu is missing", "Check the signed-in role. CareChart hides functions that are outside the role's access."),
        ("Patient is not found", "Confirm the active hospital and search using another approved identifier. Do not create a duplicate without checking."),
        ("Doctor is missing from scheduling", "Confirm an active doctor account is linked to a practitioner profile and clinic."),
        ("No appointment slots appear", "Check weekday schedule, slot length, breaks, blocked time, holidays, and existing appointments."),
        ("Same weekday appears twice", "Save the same practitioner and weekday again to update it. Ask support to review older duplicate data if it remains."),
        ("Password is rejected", "Use at least 8 characters with at least one letter and one number."),
        ("Invoice price is zero", "Create or correct the active service and clinic price in Settings, then follow the approved invoice-correction process."),
        ("Report cannot be uploaded", "Confirm the order is in progress and the file is an approved PDF, JPEG, or PNG within the configured size limit."),
        ("Doctor cannot acknowledge report", "Confirm the report is completed and the user is the ordering clinician or an authorized administrator."),
        ("Page remains on Loading", "Refresh once. If application assets return errors, contact the implementation/support team with the page and time."),
    ], [2700, 6660])
    h2(doc, "Information safe to send support")
    add_bullets(doc, [
        "Page name and approximate time of the issue",
        "Your role and the hospital name shown in the header",
        "The non-sensitive error message",
        "Steps that led to the issue and whether it remains after a refresh",
        "A screenshot only after patient information, credentials, tokens, emails, IDs, and internal URLs are removed",
    ])

    h1(doc, "20. Plain-language glossary", new_page=True)
    add_table(doc, ["Term", "Meaning"], [
        ("Active hospital", "The hospital workspace currently selected at the top of the screen."),
        ("Appointment type", "A named kind of visit with a standard duration."),
        ("Audit event", "A time-stamped record that an important action happened."),
        ("Clinical indication", "The reason a test or imaging procedure is requested."),
        ("Completed diagnostic order", "A final report was uploaded and is waiting for clinical review."),
        ("Department", "A care grouping inside a clinic, such as General Medicine or Pediatrics."),
        ("Encounter / visit", "The clinical record for one patient attendance or consultation."),
        ("IANA timezone", "A standard location-based timezone name, such as Asia/Kolkata."),
        ("MRN", "Medical Record Number - the patient's identifier inside one hospital."),
        ("Organization", "The hospital-level workspace that separates one customer's data from another's."),
        ("Practitioner profile", "The configuration linking a clinician account to clinic, department, and scheduling details."),
        ("Reviewed diagnostic order", "An authorized clinician opened and acknowledged the uploaded report."),
        ("Service catalog", "The hospital's list of billable services and clinic prices."),
        ("Tenant scoped", "Restricted to one hospital workspace rather than shared across hospitals."),
    ], [2200, 7160])
    add_callout(doc, "Guide maintenance", "Update this handbook whenever screens, roles, workflows, password rules, or integrations change. Recreate customer screenshots with fictional demonstration data and repeat the privacy review before distribution.", PALE_GREEN)

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.core_properties.title = "CareChart Customer Configuration and User Guide"
    doc.core_properties.subject = "Plain-language customer handbook for CareChart hospital configuration and daily use"
    doc.core_properties.author = "CareChart"
    doc.core_properties.keywords = "CareChart, customer guide, configuration, EMR, privacy-safe"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
