from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Admin\Documents\emr-app\output\docx\CareChart_Customer_User_Guide.docx")
IMAGES = {
    "admin": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-f45ba6dc-1999-4119-b972-f8b603dc6c28.png"),
    "settings": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-395f0c9e-f04f-43c0-a2af-1161103b79b9.png"),
    "schedule": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-2b158d6c-b98f-4939-a724-8574e4c43ed4.png"),
    "queue": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-db96ad77-84c6-4eaa-8cbe-71495c7fe9e4.png"),
    "consult": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-d0c8969a-e41a-48b4-a45d-9b0a7fbdf121.png"),
    "diagnostic": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-0aa472bb-26e7-47c2-8c43-0a428d98e2f5.png"),
    "invoice": Path(r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-7d6523d6-1bce-4a63-8867-febb13b68099.png"),
}

NAVY = "17364D"
GREEN = "3E7465"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B07A24"
INK = "17202A"
GRAY = "5C6872"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F2EE"
LIGHT_GOLD = "FFF4DD"
LIGHT_RED = "FBE9E4"
BORDER = "D9DEE5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 9, GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])


def add_numbering_definitions(doc):
    # python-docx ships stable built-in definitions: List Bullet uses numId 1
    # and List Number uses numId 5. Reusing those definitions prevents Word
    # from remapping custom list definitions during PDF export.
    return 1, 5


def set_list_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def restart_numbering(doc, base_num_id):
    """Create a fresh list instance so each procedure starts at step 1."""
    numbering = doc.part.numbering_part.element
    base = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_id = max(
        [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
        or [0]
    ) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(11)
    bullet.font.color.rgb = RGBColor.from_string(INK)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False
    if "Callout Text" not in styles:
        callout = styles.add_style("Callout Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout Text"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.space_before = Pt(0)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.2


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("CareChart Customer Guide")
    set_run_font(r, 9, GRAY, True)
    r = p.add_run("  |  Customer Edition - August 2026")
    set_run_font(r, 9, GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CareChart - Customer User and Configuration Guide")
    set_run_font(r, 8.5, GRAY)


def add_title_block(doc):
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("CUSTOMER ENABLEMENT GUIDE")
    set_run_font(r, 10, GOLD, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CareChart")
    set_run_font(r, 31, NAVY, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Complete User, Administration, and Hospital Configuration Handbook")
    set_run_font(r, 13.5, GRAY)
    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [1500, 3180, 1500, 3180])
    entries = [
        ("Prepared for", "[Hospital / Clinic Name]", "Edition", "Customer Guide"),
        ("Product", "CareChart EMR", "Updated", "25 August 2026"),
    ]
    for row, values in zip(table.rows, entries):
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, 9.5, GRAY if index % 2 == 0 else INK, bold=index % 2 == 0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("What this guide does")
    set_run_font(r, 12, GREEN, True)
    p = doc.add_paragraph("This handbook explains CareChart in everyday language. It shows hospital owners how to configure the system, staff how to complete their daily work, and customers what must be prepared before using the application with real patients.")
    p.paragraph_format.space_after = Pt(10)
    add_callout(doc, "Example data", "Names, phone numbers, clinical details, and prices shown in screenshots are fictional demonstration data. Replace them with your hospital's approved information.", LIGHT_GOLD)
    doc.add_page_break()


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout Text"]
    r = p.add_run(f"{label}: ")
    set_run_font(r, 10.5, NAVY, True)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_bullets(doc, items, numbered=False):
    base_num_id = doc._number_ids[1 if numbered else 0]
    num_id = restart_numbering(doc, base_num_id) if numbered else None
    for item in items:
        p = doc.add_paragraph()
        if numbered:
            set_list_numbering(p, num_id)
        else:
            p.style = doc.styles["List Bullet"]
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(label)
            set_run_font(r, 11, INK, True)
            r = p.add_run(detail)
            set_run_font(r, 11, INK)
        else:
            p.add_run(item)


def add_steps(doc, steps):
    add_bullets(doc, steps, numbered=True)


def add_table(doc, headers, rows, widths, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    header_pr = header._tr.get_or_add_trPr()
    header_pr.append(OxmlElement("w:cantSplit"))
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if compact:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        set_run_font(r, 9 if compact else 9.5, NAVY, True)
    for values in rows:
        row = table.add_row()
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (compact and index > 0) else WD_ALIGN_PARAGRAPH.LEFT
            if compact:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, 8.7 if compact else 9.3, INK)
    set_table_geometry(table, widths)
    return table


def add_image(doc, path, caption, alt, width=6.4):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(doc, text, new_page=False):
    heading = doc.add_heading(text, level=1)
    if new_page:
        heading.paragraph_format.page_break_before = True
    return heading


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_toc(doc):
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph("[[TOC]]")
    p.paragraph_format.space_after = Pt(12)
    add_callout(doc, "How to use this handbook", "Start with Part 1 if you are setting up a hospital. Staff members can go directly to the section for their role. The final checklists are useful during training and go-live.", LIGHT_GREEN)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    doc._number_ids = add_numbering_definitions(doc)
    add_title_block(doc)
    add_toc(doc)

    add_h1(doc, "1. Understanding CareChart")
    doc.add_paragraph("CareChart is an outpatient electronic medical record (EMR) and clinic operations application. It brings registration, appointments, consultations, diagnostic orders, reports, documents, billing, and patient access into one organization-controlled workspace.")
    add_h2(doc, "The hospital workspace idea")
    doc.add_paragraph("Every hospital is treated as a separate workspace. A normal staff member belongs to one hospital and only sees that hospital's patients, doctors, appointments, documents, invoices, and settings. A Super Admin may manage several hospitals and chooses the active hospital from the selector at the top of the screen.")
    add_callout(doc, "Important", "Always check the hospital name in the upper-left corner before entering or reviewing information. It tells you which hospital workspace is currently active.", LIGHT_GOLD)
    add_h2(doc, "Main areas")
    add_table(doc, ["Area", "What it is used for"], [
        ("Front Desk", "Register patients, search existing records, record allergies, and place patients in the doctor's queue."),
        ("Schedule", "Book appointments, review the day, manage online requests, and check patients in."),
        ("Doctor's Desk", "Review requests and the waiting room, document consultations, prescribe, order tests, and sign visits."),
        ("Patient Records", "Search the patient history and review the longitudinal clinical record."),
        ("Diagnostics", "Start laboratory or imaging work, upload the final report, and send it to the ordering doctor."),
        ("Documents", "Securely upload and retrieve PDF/JPEG/PNG patient documents."),
        ("Billing", "Generate invoices from completed visits, collect payments, record refunds, and print invoices."),
        ("Admin", "Create and manage staff, roles, invitations, and hospital accounts where permitted."),
        ("Settings", "Configure locations, doctors, schedules, appointment types, clinic prices, and tax rules."),
        ("Audit", "Review a time-stamped history of important security and workflow events."),
    ], [1800, 7560])
    add_h2(doc, "What CareChart does not do automatically")
    add_bullets(doc, [
        "It does not decide a diagnosis, treatment, medicine, dosage, or clinical priority. Qualified staff remain responsible for clinical decisions.",
        "It does not invent prices. A service without a configured clinic price stays at zero and is clearly flagged on the invoice.",
        "It does not make a development installation compliant by itself. Production hosting, backups, legal review, security operations, and approved policies are still required.",
        "External PACS/DICOM, ABDM, payment gateway, and production messaging connections require separate integration work unless your deployment partner has enabled them.",
    ])

    add_h1(doc, "2. Roles and access", new_page=True)
    doc.add_paragraph("CareChart shows each person only the functions needed for their job. This reduces confusion and lowers the chance of someone opening an area they do not use.")
    add_table(doc, ["Role", "Typical access"], [
        ("Super Admin", "Creates and switches hospitals; can configure and administer the active hospital."),
        ("Hospital Admin", "Manages users, settings, appointments, records, diagnostics, billing, documents, and audit for one hospital."),
        ("Doctor", "Uses schedule, doctor's desk, patient records, diagnostics, and documents; completes and signs clinical visits."),
        ("Nurse", "Uses appointments, patient records, and documents; supports patient and encounter preparation."),
        ("Front Desk", "Registers patients, manages appointments, opens basic records, and uploads documents."),
        ("Billing", "Reads patient details needed for billing and manages invoices, payments, refunds, and printing."),
        ("Lab / Radiology", "Reads relevant patient/order information, starts diagnostic work, and uploads reports."),
        ("Patient", "Uses the separate patient portal to view their own patient jacket and visit information."),
    ], [1900, 7460], compact=True)
    add_h2(doc, "Access rules customers should know")
    add_bullets(doc, [
        "A hospital user's organization comes from the signed-in account; staff cannot type another hospital ID into a form to change access.",
        "A suspended or removed user loses access when the server checks the account, without waiting for an old session to expire.",
        "A doctor may clinically acknowledge a report they ordered. An authorized administrator may also acknowledge when hospital policy allows it.",
        "Super Admin hospital switching changes the full active context. Patient and financial data are never combined into one cross-hospital screen.",
    ])

    add_h1(doc, "3. Signing in and password basics", new_page=True)
    add_h2(doc, "First sign-in")
    add_steps(doc, [
        "Open the web address supplied by your CareChart administrator.",
        "Enter the email address assigned to your staff account.",
        "Enter the temporary password supplied by the administrator or use the invitation link sent to your email.",
        "If CareChart asks you to change the password, create your own password before continuing.",
        "Confirm the hospital name and your role in the top navigation area.",
    ])
    add_h2(doc, "Password rule")
    doc.add_paragraph("A password must contain at least 8 characters, including at least one letter and one number. Use a unique password that is not shared with another person.")
    add_h2(doc, "Forgotten password")
    add_steps(doc, [
        "Choose Forgot password on the sign-in page.",
        "Enter the account email address.",
        "Open the time-limited reset link sent to that email.",
        "Set and confirm a new password that follows the rule above.",
    ])
    add_callout(doc, "Security reminder", "Never share passwords, reset links, patient portal activation links, or screenshots containing real patient information over an unapproved channel.", LIGHT_RED)

    add_h1(doc, "4. Super Admin: create and switch hospitals", new_page=True)
    doc.add_paragraph("This section is only for a Super Admin who manages more than one customer hospital.")
    add_h2(doc, "Create a hospital account")
    add_steps(doc, [
        "Open Admin and locate Create hospital account.",
        "Enter the hospital or clinic name.",
        "Enter the first administrator's full name and email address.",
        "Create a temporary password with at least 8 characters, a letter, and a number.",
        "Select Create hospital. CareChart creates the workspace and switches into it so configuration can begin.",
    ])
    add_callout(doc, "Super Admin membership", "The Super Admin can configure a newly created hospital immediately. The administrator email is not copied from a previous form and must be entered deliberately.", LIGHT_GREEN)
    add_image(doc, IMAGES["admin"], "Figure 1. Super Admin hospital creation and hospital-scoped administration.", "CareChart Admin screen showing the Create Hospital Account form and hospital-specific navigation.")
    add_h2(doc, "Switch hospitals")
    add_steps(doc, [
        "Find Active Hospital near the top-right corner.",
        "Open the hospital selector.",
        "Choose the hospital you want to manage.",
        "Wait for the page to refresh and verify the chosen hospital name in both the selector and upper-left heading.",
    ])
    add_callout(doc, "Before making changes", "Switching hospitals changes the patients, users, schedules, documents, prices, and settings you can see. Verify the active hospital before every configuration session.", LIGHT_GOLD)

    add_h1(doc, "5. Hospital configuration - recommended order", new_page=True)
    doc.add_paragraph("Complete configuration in the order below. Later steps depend on earlier ones. For example, a doctor schedule needs a clinic and a practitioner profile, and invoice pricing needs a clinic and service catalog.")
    add_steps(doc, [
        "Create the clinic location or campus.",
        "Add departments.",
        "Link doctor accounts to practitioner profiles.",
        "Create appointment types.",
        "Set weekly doctor schedules and optional breaks.",
        "Add blocked periods and holidays when required.",
        "Configure services, clinic prices, and taxes.",
        "Invite or create the remaining staff accounts.",
        "Run a complete test journey using fictional patient data.",
    ])
    add_image(doc, IMAGES["settings"], "Figure 2. Settings opens in the active hospital and begins with clinic/location configuration.", "CareChart Clinic Configuration screen showing the active hospital selector and fields for clinic name, code, address, and timezone.")

    add_h2(doc, "5.1 Clinic locations")
    doc.add_paragraph("A clinic location represents a physical campus, branch, or care location. Enter a clear name, a short unique code, the address, and the IANA timezone. For India, Asia/Kolkata is normally appropriate.")
    add_table(doc, ["Field", "Plain-language guidance"], [
        ("Clinic name", "Use the name staff and patients recognize, such as Main Campus or Baner Clinic."),
        ("Clinic code", "Use a short unique code, such as NLH01. Do not reuse the same code inside one hospital."),
        ("Address", "Enter the full service address used for appointment and customer information."),
        ("Timezone", "Controls how local appointment times are displayed. Database times are stored safely in UTC."),
    ], [1900, 7460])
    add_h2(doc, "5.2 Departments")
    doc.add_paragraph("Departments group practitioners within a clinic. Examples include General Medicine, Pediatrics, Cardiology, and Radiology. Select the clinic, choose Department as the configuration type, enter the department name, and save.")
    add_h2(doc, "5.3 Practitioner profiles")
    doc.add_paragraph("A staff account with the Doctor role becomes schedulable after it is linked to a practitioner profile. Select the doctor, clinic, department, specialty, qualification, registration number, and default appointment length. Use the doctor's official registration details.")
    add_h2(doc, "5.4 Appointment types")
    doc.add_paragraph("Appointment types tell CareChart what kind of visit can be booked and how long it normally takes. Examples: General Consultation - 30 minutes, Follow-up - 15 minutes, Pediatric Consultation - 30 minutes.")
    add_h2(doc, "5.5 Weekly schedules")
    add_steps(doc, [
        "Select the practitioner and weekday.",
        "Enter the start and end time.",
        "Enter the slot length in minutes.",
        "Use Add break only when the practitioner has a break. Leave it removed when there is no break.",
        "Select Save weekly schedule.",
        "Repeat for each working day.",
    ])
    add_callout(doc, "Updating a day", "Saving the same practitioner and weekday again overwrites that day's existing schedule. It does not create a duplicate Saturday, Monday, or other weekday entry.", LIGHT_GREEN)
    add_image(doc, IMAGES["schedule"], "Figure 3. Weekly schedule fields for practitioner, day, hours, slot length, and an optional break.", "CareChart schedule configuration showing Dr. Ananya Sharma, Saturday working hours, slot minutes, and optional break controls.")
    add_h2(doc, "5.6 Blocked periods and holidays")
    doc.add_paragraph("Use blocked periods for doctor leave, training, or a temporary unavailable window. Use holidays when a whole clinic is closed. These controls prevent those times from being offered as available appointment slots.")
    add_h2(doc, "5.7 Service catalog, clinic prices, and tax")
    doc.add_paragraph("The service catalog connects the care delivered to invoice prices. Create one entry per billable service and assign a price for each clinic. Use clear names that match what staff enter during care, such as General Consultation, Complete Blood Count, Paracetamol, or Chest X-ray.")
    add_table(doc, ["Field", "Example", "Why it matters"], [
        ("Code", "CONS_GEN", "A short unique internal identifier."),
        ("Name", "General Consultation", "Used to match the visit line to the configured price."),
        ("Category", "Consultation", "Keeps consultation, medicine, test, imaging, and other charges separate."),
        ("Clinic price", "800.00", "The unit amount automatically copied into a new invoice."),
        ("Taxable", "No", "Choose according to approved billing and tax guidance."),
    ], [1500, 1700, 6160])
    add_callout(doc, "Pricing rule", "CareChart only auto-fills a price that has been configured for the visit's clinic. Missing prices remain zero and are highlighted before payment is collected.", LIGHT_GOLD)

    add_h1(doc, "6. Staff administration", new_page=True)
    add_h2(doc, "Invite staff by email")
    add_steps(doc, [
        "Open Admin in the correct hospital workspace.",
        "Enter the staff member's work email.",
        "Choose the role that matches their actual job.",
        "Select Send invite.",
        "The staff member opens the link, creates their password, and joins that hospital.",
    ])
    add_h2(doc, "Create an account manually")
    add_steps(doc, [
        "Open Add user manually.",
        "Enter the full name and work email.",
        "Create a temporary password with at least 8 characters, a letter, and a number.",
        "Choose the correct role and select Create account.",
        "Send the temporary password using your organization's approved secure process. The user is asked to change it at first sign-in.",
    ])
    add_h2(doc, "Manage users")
    add_bullets(doc, [
        "Approve or reject pending join requests.",
        "Change an active user's role when their job changes.",
        "Revoke an invitation that should no longer be used.",
        "Remove a staff member when they leave the hospital; access is lost immediately.",
        "Do not share a generic login between several people. Individual accounts make the audit history meaningful.",
    ])

    add_h1(doc, "7. Front Desk workflow", new_page=True)
    add_h2(doc, "Find an existing patient first")
    doc.add_paragraph("Search by name, medical record number (MRN), or phone number before registering anyone. This reduces duplicate patient files.")
    add_h2(doc, "Register a new patient")
    add_steps(doc, [
        "Open Front Desk and select New Patient Registration.",
        "Enter the patient's name, age/date of birth, gender, phone, email, address, blood group, and emergency details where available.",
        "Record known allergies exactly as reported and verify them with clinical staff.",
        "Save the patient. CareChart assigns an MRN inside the active hospital.",
        "If a possible duplicate appears, stop and ask an administrator to review it instead of creating another record.",
    ])
    add_h2(doc, "Log a walk-in visit")
    add_steps(doc, [
        "Open the patient's record from Front Desk.",
        "Choose the doctor and record the chief complaint and available basic observations.",
        "Log the visit. The patient appears in that doctor's waiting room.",
        "Tell the clinical team if an allergy or urgent concern was reported.",
    ])
    add_callout(doc, "Patient identity", "Do not merge two records just because their names look similar. Administrators should verify date of birth, phone, identifiers, and clinical history before an audited merge.", LIGHT_RED)

    add_h1(doc, "8. Appointments and online booking", new_page=True)
    add_h2(doc, "Book an appointment for an existing patient")
    add_steps(doc, [
        "Open Schedule and choose the date.",
        "Select New appointment.",
        "Search for and select the patient.",
        "Choose a doctor. CareChart shows the doctor's clinic and specialty.",
        "Choose an appointment type, date/time, and reason.",
        "Save. CareChart checks the selected time against configured availability and existing reservations.",
    ])
    add_h2(doc, "Public online requests")
    doc.add_paragraph("A patient can use the public appointment page to choose an available doctor and time, enter contact details and a reason, accept the privacy statement, and submit a request. The time is reserved while the doctor reviews it. The patient receives a booking reference and, when email is configured, a confirmation message after approval.")
    add_h2(doc, "Check in")
    add_steps(doc, [
        "Open the appointment on the day of arrival.",
        "For an online patient, complete the required patient record details such as age and gender.",
        "Select Check in. The appointment moves into the doctor's queue.",
    ])
    add_table(doc, ["Status", "Meaning"], [
        ("Requested", "The patient asked for the time; a doctor or authorized user must review it."),
        ("Confirmed / Scheduled", "The appointment is accepted and on the calendar."),
        ("Checked in / In consultation", "The patient has arrived; the doctor may begin and continue the encounter."),
        ("Completed", "The consultation is finished and may be billed."),
        ("Cancelled / No show / Rescheduled", "The original appointment did not proceed as originally planned."),
    ], [1900, 7460])

    add_h1(doc, "9. Doctor's Desk and consultation", new_page=True)
    doc.add_paragraph("The Doctor's Desk brings together online requests, confirmed appointments, and the walk-in waiting room.")
    add_image(doc, IMAGES["queue"], "Figure 4. Doctor's Desk with appointment-request, confirmed-queue, and waiting-room sections.", "CareChart Doctor's Desk showing a waiting-room patient and no pending online requests.")
    add_h2(doc, "Review requests")
    add_bullets(doc, [
        "Approve a suitable online appointment request to confirm it.",
        "Reject or cancel only with an appropriate reason when required by hospital policy.",
        "Open a waiting-room patient to begin the consultation.",
    ])
    add_h2(doc, "Complete the consultation")
    add_steps(doc, [
        "Confirm the patient's identity and read the allergy alert before entering clinical information.",
        "Review the chief complaint, vitals, and relevant visit history.",
        "Record examination findings, diagnosis, and advice in clear language.",
        "Add medicines one at a time with dosage, frequency, and duration.",
        "Review every allergy, duplicate-therapy, or interaction warning. Record an acknowledgement or clinical reason only when justified.",
        "Add requested tests as separate items so billing and follow-up can match them correctly.",
        "Create an operational diagnostic order when the lab/radiology team must receive and track the request.",
        "Save a draft if the visit is not finished. Otherwise confirm the signature checkbox and select Complete & sign visit. After signing, use Print report when a patient or approved workflow needs a formatted consultation report.",
    ])
    add_image(doc, IMAGES["consult"], "Figure 5. Consultation form with allergy alert, complaint, vitals, examination, diagnosis, advice, medicines, and orders.", "CareChart consultation screen for a fictional patient showing allergy alert and clinical documentation fields.")
    add_h2(doc, "Draft versus signed visit")
    add_table(doc, ["Action", "Use it when"], [
        ("Save draft", "Information is still being collected or the clinician is not ready to finalize."),
        ("Complete & sign visit", "The clinician has checked the record and is ready to make it the finalized encounter."),
        ("Amendment", "A signed encounter needs a documented correction. The original content is retained for traceability."),
    ], [2100, 7260])
    add_callout(doc, "Clinical responsibility", "CareChart records and warns; it does not replace professional judgement. The signing clinician remains responsible for accuracy and appropriateness.", LIGHT_RED)

    # Let this chapter follow the short signed-visit continuation so the
    # customer edition does not contain a mostly empty page.
    add_h1(doc, "10. Diagnostic orders and report review")
    add_h2(doc, "Doctor creates the order")
    add_steps(doc, [
        "Inside the consultation, choose Laboratory or Imaging.",
        "Choose Routine, Urgent, or STAT according to clinical judgement and hospital policy.",
        "Enter the procedure name and optional code.",
        "Enter the clinical indication and an optional scheduled time.",
        "Select Create diagnostic order.",
    ])
    add_h2(doc, "Lab/radiology processes the order")
    add_steps(doc, [
        "Open Diagnostics and find the order by type or status.",
        "For a blood test, select Collect sample & start when the patient gives the sample. For imaging, select Start processing.",
        "Complete the laboratory or imaging work outside CareChart using the hospital's approved process.",
        "When the final result is ready, upload the PDF, JPEG, or PNG directly inside the in-progress order.",
        "CareChart scans and stores the file, links it to the order, and changes the order to Completed (ready for review).",
    ])
    add_image(doc, IMAGES["diagnostic"], "Figure 6. In-progress diagnostic order ready for final report upload and doctor review.", "CareChart diagnostics screen showing an in-progress laboratory order for a fictional patient.")
    add_h2(doc, "Doctor reviews the report")
    add_steps(doc, [
        "Open the completed order from Diagnostics.",
        "Select Open report and review the attached file.",
        "If clinically satisfied, select Review & acknowledge and confirm the prompt.",
        "The order changes to Reviewed and the action is recorded in the audit history.",
    ])
    add_callout(doc, "Do not use the status as a substitute for reading", "Completed means the report was uploaded. Reviewed means an authorized clinician actually opened, reviewed, and acknowledged it.", LIGHT_GOLD)

    add_h1(doc, "11. Patient Records and patient identity", new_page=True)
    add_h2(doc, "Search and review")
    doc.add_paragraph("Patient Records provides a searchable history for the active hospital. Staff with permission can review demographics, allergies, past encounters, diagnoses, prescriptions, tests, documents, and signed visit details.")
    add_h2(doc, "Longitudinal clinical summary")
    add_bullets(doc, [
        "Active problems and diagnoses",
        "Current medication statements",
        "Detailed allergies and reactions",
        "Immunizations",
        "Procedures",
        "Safety or clinical flags",
        "Reviewed diagnostic observations",
        "Past visits and signed encounter information",
    ])
    add_h2(doc, "Additional identifiers")
    doc.add_paragraph("Administrators may record additional identifiers when needed, such as an external hospital number or approved national/partner identifier. The system keeps identifiers scoped to the hospital and avoids reusing the same system/value combination.")
    add_h2(doc, "Duplicate review and merge")
    add_steps(doc, [
        "Search possible duplicates using normalized phone, date of birth, name, and other strong details.",
        "Open both records and verify that they represent the same person.",
        "Choose the surviving record carefully.",
        "Enter a clear merge reason and confirm the administrator-only merge.",
        "CareChart records the source snapshot and audit event. A merge is not casual and should follow hospital policy.",
    ])
    add_h2(doc, "FHIR-style export")
    doc.add_paragraph("Authorized staff can obtain an audited, hospital-scoped FHIR R4-style patient bundle for supported information-exchange workflows. This is a technical export and should be used only with an approved receiving system and privacy process.")

    # Continue after the short export subsection to avoid a nearly empty page.
    add_h1(doc, "12. Patient documents")
    add_steps(doc, [
        "Open Documents and search for the patient.",
        "Choose a PDF, JPEG, or PNG file.",
        "Optionally link it to a diagnostic order; otherwise it remains a general patient document.",
        "Select Upload securely.",
        "Authorized users can later open or download the document from the same hospital workspace.",
    ])
    add_h2(doc, "Protection applied by the application")
    add_bullets(doc, [
        "Only approved file types are accepted and the file content must match its declared type.",
        "A configurable size limit is enforced (10 MB by default in the supplied environment example).",
        "Files are checked by the configured malware scanner. Production must use a real scanner endpoint.",
        "Random private storage keys are used instead of patient names in server paths.",
        "Checksums detect unexpected file changes, and production storage encryption requires a configured key.",
        "Every download checks the signed-in user's active hospital and creates an audit event.",
    ])

    add_h1(doc, "13. Billing, invoices, and payments", new_page=True)
    add_h2(doc, "Generate an invoice")
    add_steps(doc, [
        "Complete and sign the clinical visit.",
        "Open Billing. The visit appears under Completed visits awaiting invoice.",
        "Select Generate invoice.",
        "CareChart identifies the visit's clinic and matches consultation, medicine, test, and imaging lines to that clinic's configured catalog prices.",
        "Review all lines, quantities, discounts, tax, and the grand total before saving or collecting money.",
    ])
    add_image(doc, IMAGES["invoice"], "Figure 7. Generated invoice with editable line items, totals, payment area, printing, and void controls.", "CareChart invoice screen showing consultation, medicine, and test line items for a fictional patient.")
    add_h2(doc, "Automatic pricing behavior")
    add_bullets(doc, [
        "General Consultation matches the configured consultation service for that clinic.",
        "A medicine matches its service catalog name even when the prescription also shows a strength or dosage.",
        "Comma-separated tests become separate invoice lines so each test can be priced independently.",
        "Tax is applied only when the matched service is marked taxable and an active clinic tax rule exists.",
        "If a match is missing, the unit price remains zero and CareChart lists the unconfigured services in a warning.",
    ])
    add_h2(doc, "Payments, refunds, and voids")
    add_table(doc, ["Action", "Plain-language meaning"], [
        ("Record payment", "Enter the amount, payment method, and optional reference. CareChart tracks remaining balance and produces receipt information."),
        ("Partial payment", "The invoice remains partly unpaid until the balance is collected."),
        ("Refund", "Record a refund amount and reason separately. The original payment remains visible for financial traceability."),
        ("Void invoice", "Use only when the invoice should be cancelled. A reason is required and the action cannot be treated as an ordinary edit."),
        ("Print", "Print or save a customer-facing copy using the browser's print function."),
    ], [1800, 7560])
    add_callout(doc, "Before payment", "Never collect against a zero-priced line unless it is intentionally free and approved. Configure missing services under Settings - Billing, then follow the hospital's correction process for an already-created invoice.", LIGHT_RED)

    add_h1(doc, "14. Patient portal", new_page=True)
    add_h2(doc, "Activate access")
    add_steps(doc, [
        "The hospital creates or supplies a patient portal activation link using its approved process.",
        "The patient opens the link, verifies the requested information, and creates a private password.",
        "The patient signs in through the separate Patient Portal page, not the staff login page.",
    ])
    add_h2(doc, "What the patient can see")
    add_bullets(doc, [
        "Name, MRN, age, gender, hospital, and blood group",
        "The next appointment when one exists",
        "Counts and summaries for allergies, visits, and medicines",
        "Past visit details, diagnosis, vitals, doctor's notes, advice, prescriptions, tests, and digital signature where available",
    ])
    add_callout(doc, "Patient privacy", "A portal user sees only the patient jacket connected to that portal account. Staff should verify the correct email and identity before activation.", LIGHT_GREEN)

    add_h1(doc, "15. Audit, privacy, and safety controls", new_page=True)
    add_h2(doc, "Audit history")
    doc.add_paragraph("Audit is an append-only list of important actions for the active hospital. Administrators can review the time, action, resource, user/system identifier, and reason where one was supplied.")
    add_bullets(doc, [
        "Sign-ins, failed access attempts, and password events",
        "Hospital switching and administration changes",
        "Patient creation, updates, identity merges, and exports",
        "Appointment and encounter workflow changes",
        "Diagnostic order, report, and review actions",
        "Document uploads, downloads, and integrity failures",
        "Invoices, payments, refunds, and voids",
    ])
    add_h2(doc, "Consent and privacy requests")
    doc.add_paragraph("CareChart includes consent and privacy-request records so the hospital can document requests such as access, correction, export, erasure review, or grievance. These records support a process; they do not automatically delete clinical or financial records.")
    add_h2(doc, "Hospital responsibilities")
    add_bullets(doc, [
        "Give access only to people who need it for their work.",
        "Review active staff regularly and remove leavers promptly.",
        "Use HTTPS, strong secrets, encrypted backups, and approved production hosting.",
        "Define retention, correction, disclosure, and breach-response policies.",
        "Train staff not to put passwords, tokens, or unnecessary clinical text into audit reasons or support messages.",
        "Complete legal and regulatory review for the country and care setting, including applicable Indian DPDP and ABDM guidance when relevant.",
    ])

    add_h1(doc, "16. End-to-end daily workflows", new_page=True)
    add_h2(doc, "Walk-in consultation")
    add_steps(doc, [
        "Front Desk searches or registers the patient.",
        "Front Desk logs the visit and chooses the doctor.",
        "Doctor opens the patient from the waiting room.",
        "Doctor documents, prescribes/orders, and signs the visit.",
        "Lab/radiology completes any tracked diagnostic order and uploads the report.",
        "Ordering doctor reviews and acknowledges the report.",
        "Billing generates the clinic-priced invoice and records payment.",
        "Authorized staff can later retrieve the signed visit, report, invoice, and audit history.",
    ])
    add_h2(doc, "Online appointment")
    add_steps(doc, [
        "Patient requests a doctor, date, and open time online.",
        "Doctor reviews and confirms the request.",
        "Front Desk checks the patient in and completes any missing demographics.",
        "The consultation, diagnostics, records, and billing follow the same controlled workflow as a walk-in visit.",
    ])
    add_h2(doc, "New hospital setup checklist")
    add_bullets(doc, [
        "Hospital workspace created and active hospital verified",
        "Clinic locations, addresses, codes, and timezones entered",
        "Departments and appointment types created",
        "Doctor accounts linked to practitioner profiles",
        "Every working day and optional break checked",
        "Holidays and planned leave entered",
        "Common consultations, medicines, tests, and imaging prices configured per clinic",
        "Tax rules reviewed by an authorized finance professional",
        "Staff invited with least-privilege roles",
        "Email, secure storage, encryption, scanning, backups, and HTTPS tested",
        "Fictional end-to-end test completed before real patient entry",
    ])

    add_h1(doc, "17. Troubleshooting", new_page=True)
    add_table(doc, ["Problem", "What to check"], [
        ("Wrong hospital name is shown", "A Super Admin should open Active Hospital and switch. Other users must be assigned to the correct hospital account."),
        ("A staff member cannot see a menu", "Check their role. CareChart intentionally hides functions outside that role."),
        ("Doctor is missing from scheduling", "Confirm an active Doctor account has a practitioner profile linked to a clinic."),
        ("No appointment times appear", "Check the doctor's weekday schedule, slot length, breaks, blocked periods, holidays, and existing bookings."),
        ("The same weekday appears twice", "Save the same practitioner/day again; the current schedule overwrites the older entry. Contact support if old duplicate data remains."),
        ("Password is rejected", "Use at least 8 characters with at least one letter and one number."),
        ("Invoice price is zero", "Create an active service catalog price for that exact service/category in the visit's clinic."),
        ("Report cannot be uploaded", "Confirm the order is In progress and the file is a valid PDF/JPEG/PNG within the configured size limit."),
        ("Doctor cannot acknowledge a report", "Confirm the order is Completed and the signed-in doctor is the ordering doctor, or use an authorized administrator."),
        ("Page stays on Loading", "Refresh once. If JavaScript/CSS files return 404, restart the application server and regenerate the current build assets."),
        ("Patient appears twice", "Do not delete casually. Ask an administrator to verify and use the audited duplicate-merge workflow."),
    ], [2500, 6860])
    add_h2(doc, "Information to send support")
    add_bullets(doc, [
        "The page name and approximate time the issue happened",
        "Your role and the hospital name shown in the header",
        "The safe error message (do not include a password, reset link, or secret)",
        "An approved screenshot with patient information hidden when possible",
        "The steps that led to the issue and whether it happens after a refresh",
    ])

    add_h1(doc, "18. Production readiness and current limits", new_page=True)
    add_callout(doc, "Customer disclosure", "CareChart is a strong outpatient workflow foundation, but the supplied development application is not a legal, security, or regulatory certification. A production launch requires technical, operational, clinical, and legal readiness work.", LIGHT_RED)
    add_h2(doc, "Minimum production preparation")
    add_bullets(doc, [
        "Deploy behind HTTPS using a managed production environment.",
        "Use a strong random NextAuth secret and protected production database credentials.",
        "Use managed PostgreSQL backups, test restoration, and approve a retention schedule.",
        "Configure a verified email sender for appointment and password communications.",
        "Use production object storage, a 32-byte document encryption key, and a real malware-scanner service.",
        "Review permissions, audit monitoring, onboarding/offboarding, incident response, and support access.",
        "Complete clinical safety testing, user acceptance testing, security testing, and applicable privacy/regulatory review.",
    ])
    add_h2(doc, "Integration status")
    add_table(doc, ["Integration", "Current guide position"], [
        ("Email", "Supported when Resend and a verified sender are configured."),
        ("HL7 imaging message", "Local workflow/adaptor foundation exists; production transport and receiving-system validation are deployment work."),
        ("PACS / DICOM", "Not enabled as a complete production integration in the supplied application."),
        ("ABDM", "Not enabled as a complete production integration; use only after separate design, consent, and conformance work."),
        ("Payment gateway", "Not connected; staff record payments received through approved external methods."),
        ("FHIR-style export", "A tenant-scoped patient export is available for approved technical workflows, not automatic exchange."),
    ], [2200, 7160])

    add_h1(doc, "19. Plain-language glossary", new_page=True)
    add_table(doc, ["Term", "Meaning"], [
        ("Active hospital", "The hospital workspace currently selected. All data and configuration on screen belong to it."),
        ("Appointment type", "A named kind of visit with a standard length, such as General Consultation - 30 minutes."),
        ("Audit event", "A time-stamped record showing that an important action happened."),
        ("Clinical indication", "Why a test or imaging procedure is being requested."),
        ("Completed diagnostic order", "The final report has been uploaded and is waiting for clinical review."),
        ("Department", "A clinical grouping inside a location, such as Pediatrics or General Medicine."),
        ("Encounter / visit", "The clinical record for one patient attendance or consultation."),
        ("FHIR", "A standard structure used to exchange health information between approved systems."),
        ("IANA timezone", "A standard location-based timezone name, such as Asia/Kolkata."),
        ("MRN", "Medical Record Number - the patient's identifier inside a hospital."),
        ("Organization", "The hospital-level workspace that keeps one customer's data separate from another's."),
        ("Practitioner profile", "The scheduling and clinical details that connect a doctor account to a clinic and department."),
        ("Reviewed diagnostic order", "An authorized clinician has opened and acknowledged the uploaded report."),
        ("Service catalog", "The hospital's list of billable services and the names used to match invoice lines."),
        ("Tenant scoped", "Restricted to one hospital workspace rather than shared across all hospitals."),
    ], [1900, 7460])
    add_callout(doc, "Guide maintenance", "Update this handbook when workflows, role permissions, integrations, password rules, or customer-facing screens change. Put the revision date on the cover so customers know which edition they received.", LIGHT_GREEN)

    # Ask Word/LibreOffice to refresh fields such as TOC/page numbers when opened.
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.core_properties.title = "CareChart Customer User and Configuration Guide"
    doc.core_properties.subject = "Customer-facing CareChart EMR handbook"
    doc.core_properties.author = "CareChart"
    doc.core_properties.keywords = "CareChart, EMR, customer guide, hospital configuration, user manual"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
