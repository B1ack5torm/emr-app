from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
SCREEN_DIR = ROOT / "screens" / "EMR App"
OUT = ROOT / "CareChart_Customer_Configuration_and_User_Guide_Professional.docx"

INK = "17202A"
GREEN = "2E6B5A"
GREEN_DARK = "204D41"
GREEN_PALE = "EAF2EE"
CREAM = "F7F4EC"
MUTED = "5C6C66"
BORDER = "D9D9D9"
PALE_BLUE = "EEF4F8"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "5")
        tag.set(qn("w:color"), BORDER)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Aptos", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=8, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.27)
section.footer_distance = Inches(0.28)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for style_name, size, before, after in (("Title", 28, 0, 12), ("Heading 1", 20, 0, 9), ("Heading 2", 13, 10, 5)):
    style = styles[style_name]
    style.font.name = "Aptos Display" if style_name != "Heading 2" else "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string("000000")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    if style_name == "Title":
        title_ppr = style._element.get_or_add_pPr()
        title_border = title_ppr.find(qn("w:pBdr"))
        if title_border is not None:
            title_ppr.remove(title_border)

if "Eyebrow" not in styles:
    eyebrow = styles.add_style("Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
else:
    eyebrow = styles["Eyebrow"]
eyebrow.font.name = "Aptos"
eyebrow._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
eyebrow._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
eyebrow.font.size = Pt(8)
eyebrow.font.bold = True
eyebrow.font.color.rgb = RGBColor.from_string(GREEN_DARK)
eyebrow.paragraph_format.space_after = Pt(4)

for sec in doc.sections:
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("CareChart  |  Customer Configuration and User Guide")
    set_font(r, size=8, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("CareChart  |  September 2026  |  ")
    set_font(r, size=8, color=MUTED)
    add_page_number(fp)


def page_break():
    doc.add_page_break()


def title_block(section_name, title, intro=None):
    p = doc.add_paragraph(style="Eyebrow")
    p.add_run(section_name.upper())
    doc.add_paragraph(title, style="Heading 1")
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(9)


def add_bullets(items, numbered=False):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        if numbered:
            prefix = p.add_run(f"{index}. ")
            prefix.bold = True
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_lead(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    p.add_run(text)


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, GREEN_DARK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=9, bold=True, color="FFFFFF")
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cell = cells[index]
            if row_index % 2:
                set_cell_shading(cell, PALE_BLUE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=8.8, color=INK)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_screenshot(filename, caption, max_height=4.45, max_width=7.05):
    path = SCREEN_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        ratio = image.width / image.height
    width = min(max_width, max_height * ratio)
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph(caption)
    cp.style = styles["Caption"]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(7)
    for run in cp.runs:
        set_font(run, size=8, color=MUTED)


def new_content_page(section_name, title, intro=None):
    page_break()
    title_block(section_name, title, intro)


# Cover
doc.core_properties.title = "CareChart Customer Configuration and User Guide"
doc.core_properties.subject = "Doctor-facing handbook for CareChart hospital setup and daily workflows"
doc.core_properties.author = "CareChart"
doc.core_properties.keywords = "CareChart, EMR, hospital, doctor, configuration, user guide"

p = doc.add_paragraph(style="Eyebrow")
p.paragraph_format.space_before = Pt(20)
p.add_run("CUSTOMER GUIDE")
doc.add_paragraph("CareChart Customer Configuration and User Guide", style="Title")
p = doc.add_paragraph("Hospital setup and daily workflows for doctors and care teams")
p.paragraph_format.space_after = Pt(14)
for run in p.runs:
    set_font(run, size=14, color=MUTED)
add_screenshot("Main Page - 01.png", "CareChart public home page and access points", max_height=3.5)
add_table(
    ["Audience", "Edition", "Scope"],
    [["Doctors, administrators and authorized hospital staff", "Customer edition - September 2026", "Outpatient setup, clinical workflow, diagnostics, billing and patient access"]],
    [2.25, 1.8, 3.0],
)
add_lead("Demonstration data", "All names, identifiers and clinical information shown in this guide are fictional and are intended only for training and presentation.")

# Overview
new_content_page("Product overview", "A connected outpatient workflow", "CareChart brings registration, appointments, consultations, diagnostics, records, billing and patient access into one hospital-scoped workspace.")
add_screenshot("Main Page - 02.png", "Product overview showing the main hospital workflows", max_height=4.0)
add_bullets([
    ("One patient record. ", "Clinical, diagnostic and billing information stays connected to the same hospital record."),
    ("Role-based access. ", "Staff see the areas needed for their assigned work."),
    ("Traceable activity. ", "Important security and operational actions are recorded for authorized review."),
])
add_lead("Clinical responsibility", "CareChart records information and displays workflow alerts. It does not replace professional judgement, hospital policy or the clinician's responsibility for patient care.")

new_content_page("Product overview", "What the platform covers", "The customer experience extends from appointment request through clinical documentation, reports, invoices and patient access.")
add_screenshot("Main Page - 03.png", "CareChart capability summary", max_height=3.95)
add_table(
    ["Area", "Primary purpose"],
    [
        ["Front Desk", "Find or register patients, manage arrivals and begin visits."],
        ["Doctor's Desk", "Open the waiting room and document consultations."],
        ["Patient Records", "Review permitted longitudinal patient information."],
        ["Diagnostics", "Process operational orders, reports and clinical acknowledgement."],
        ["Billing", "Generate invoices and record payments or refunds."],
        ["Patient Portal", "Give an activated patient access to their own information."],
    ],
    [1.7, 5.35],
)

new_content_page("Guide map", "Use the guide by role", "Administrators should complete hospital configuration first. Other staff can begin with the section that matches their daily work.")
add_table(
    ["Role", "Start with", "Typical responsibility"],
    [
        ["Hospital administrator", "Administration and configuration", "Hospital account, staff access, clinics, schedules, prices and order capabilities"],
        ["Doctor", "Doctor's Desk and consultation", "Clinical documentation, prescriptions, requested tests and imaging, report review"],
        ["Front desk", "Patient registration and appointments", "Patient identity, arrivals, appointments and queue entry"],
        ["Lab or radiology", "Diagnostics", "Operational worklist processing and report upload"],
        ["Billing", "Invoices and payments", "Invoice review, payment, refund, print and void workflows"],
        ["Patient", "Patient portal", "Account activation and access to the linked patient record"],
    ],
    [1.55, 2.15, 3.35],
)
doc.add_paragraph("Recommended reading order", style="Heading 2")
add_bullets([
    "Access and hospital selection",
    "Hospital configuration and staff setup",
    "Daily patient and clinical workflows",
    "Diagnostics, documents, billing and audit",
    "Patient portal, go-live controls and troubleshooting",
], numbered=True)
add_lead("Before entering real patient information", "Complete the hospital's clinical, privacy, legal, security, backup and operational approvals. Test the full workflow with fictional patients first.")

# Access and administration
new_content_page("Access", "Staff sign in", "Each staff member should use an individual account. Shared credentials prevent reliable access control and audit review.")
add_screenshot("Staff Sign-in.png", "Staff sign-in page", max_height=4.4)
add_bullets([
    "Enter the email address assigned to your account.",
    "Enter your password and select Sign in.",
    "Use the approved password-reset process when access is lost.",
    "Confirm the hospital shown after sign-in before opening patient information.",
], numbered=True)
add_lead("Safe access", "Never share passwords, reuse another employee's account or leave a signed-in device unattended.")

new_content_page("Administration", "Create and select a hospital workspace", "Super Admin users can create hospital accounts and switch between the hospital workspaces assigned to them.")
add_screenshot("Admin - 01.png", "Administration page for hospital creation and staff onboarding", max_height=4.15)
add_bullets([
    "Enter the hospital name and the first administrator's details.",
    "Create the hospital, then select it from Active Hospital.",
    "Verify the selected hospital before configuring clinics or inviting staff.",
], numbered=True)
add_lead("Workspace boundary", "Changing the active hospital changes the workspace. It does not move patients, visits, invoices, staff or settings between hospitals.")

new_content_page("Administration", "Review staff access", "Administrators should keep the staff list current and assign the smallest role needed for each person's work.")
add_screenshot("Admin - 02.png", "Active staff list and role management", max_height=4.2)
add_bullets([
    "Confirm the person's name, email and current duties before changing a role.",
    "Remove or suspend access promptly when a person leaves or no longer needs the system.",
    "Never reuse a former employee's account for another staff member.",
])

# Settings
new_content_page("Hospital configuration", "Configure locations and departments", "Configuration belongs to the active hospital. Begin with the locations where care is delivered.")
add_screenshot("Settings - 01.png", "Clinic locations and department configuration", max_height=4.2)
add_bullets([
    "Enter the clinic name, code, address and IANA timezone, such as Asia/Kolkata.",
    "Add the departments and appointment types used at that location.",
    "Use stable clinic codes so reporting and integrations remain consistent.",
], numbered=True)

new_content_page("Hospital configuration", "Link practitioners and weekly schedules", "A doctor account must be linked to a practitioner profile and clinic before reliable scheduling can begin.")
add_screenshot("Settings - 02.png", "Practitioner profile and weekly schedule configuration", max_height=4.15)
add_bullets([
    "Choose the doctor, clinic and department.",
    "Record specialty, qualification and registration number as required by hospital policy.",
    "Set each working weekday, start and end time, slot length and optional break.",
    "Review the saved weekly schedule before publishing appointment availability.",
], numbered=True)

new_content_page("Hospital configuration", "Control blocked time and holidays", "Use blocked periods for practitioner-specific unavailability and clinic holidays when an entire location is closed.")
add_screenshot("Settings - 03.png", "Blocked practitioner time and clinic holiday controls", max_height=3.55)
add_bullets([
    ("Blocked time. ", "Record leave, meetings or another period when one practitioner cannot accept appointments."),
    ("Clinic holiday. ", "Close appointment availability for the selected clinic and date."),
    ("Clear labels. ", "Enter a useful reason or holiday name so staff can understand why slots are unavailable."),
])

new_content_page("Hospital configuration", "Enable only the operational order systems in use", "CareChart separates clinical recommendations from operational orders sent to hospital departments or imaging systems.")
add_screenshot("Settings - 04.png", "Hospital-level controls for operational diagnostic and imaging orders", max_height=2.0)
add_bullets([
    ("Operational diagnostic orders. ", "Enable only when the hospital has a laboratory or diagnostic department that will receive and process orders."),
    ("Operational imaging orders. ", "Enable only when radiology equipment or a connected HL7 modality worklist is configured and validated."),
    ("Default state. ", "Both controls remain off until an authorized administrator enables them."),
])
add_lead("Doctor workflow", "Doctors can still record requested investigations under Tests ordered and Imaging ordered when operational integrations are disabled.")
add_screenshot("Settings - 05.png", "Service catalog, clinic pricing and tax configuration", max_height=3.05)

# Front desk and appointments
new_content_page("Front desk", "Manage arrivals from one workspace", "The Front Desk workspace combines the daily appointment schedule with patient search and walk-in registration.")
add_screenshot("Front Desk View.png", "Front Desk schedule and patient arrival panel", max_height=4.15)
add_bullets([
    "Search by MRN, patient name or phone before creating a new record.",
    "Verify more than one identifier before selecting a patient.",
    "Register a walk-in only when no correct existing record is found.",
    "Check scheduled patients in when they arrive so they enter the doctor's queue.",
], numbered=True)

new_content_page("Front desk", "Register a new patient", "Capture the minimum information required for safe identification and the current visit.")
add_screenshot("New Patient Registration.png", "New patient registration with simplified day, month and year entry", max_height=4.25)
add_bullets([
    "Enter the patient's full name as provided.",
    "Choose day and month, then type the four-digit birth year. Age is calculated automatically.",
    "If the exact birth date is unknown, leave all three birth-date fields blank and enter age.",
    "Record gender, doctor, contact information and address according to hospital policy.",
    "Enter each known allergy separately and review the information before saving.",
], numbered=True)

new_content_page("Appointments", "Book an appointment from the schedule", "Authorized staff can create an appointment for an existing patient and assign it to the appropriate doctor and time.")
add_screenshot("Manual Appointment.png", "Manual appointment form", max_height=4.75, max_width=4.6)
add_bullets([
    "Search for and select the correct patient.",
    "Choose the doctor, time, duration and reason for appointment.",
    "Review availability, then select Book appointment.",
], numbered=True)

new_content_page("Appointments", "Offer public appointment requests", "The public booking page allows a patient to request an available doctor and time without using the staff workspace.")
add_screenshot("Book Appointment.png", "Public appointment request page", max_height=4.25)
add_bullets([
    "The patient chooses the hospital, clinic, doctor, date and available time.",
    "A request is not a completed visit. Authorized staff or the doctor reviews it before confirmation.",
    "Front Desk checks the patient in on arrival.",
])

# Doctor workflow
new_content_page("Doctor workflow", "Open the waiting room", "Doctor's Desk combines appointment requests, confirmed appointments and patients waiting for consultation.")
add_screenshot("Doctor's Desk.png", "Doctor's Desk appointment queues and waiting room", max_height=4.2)
add_bullets([
    "Review pending appointment requests according to the hospital's booking process.",
    "Open a patient from Waiting Room only when ready to begin the consultation.",
    "Confirm the patient and visit reason before documenting clinical information.",
], numbered=True)

new_content_page("Doctor workflow", "Document the consultation", "The consultation begins with patient identity, allergy status, chief complaint and available vital signs.")
add_screenshot("Doctor's Desk Patient Form - 01.png", "Clinical consultation form with patient header and primary documentation fields", max_height=4.1)
add_bullets([
    "Verify the patient header and review allergy information before prescribing.",
    "Record examination findings, diagnosis and advice in clear clinical language.",
    "Add medicines only after reviewing allergies, duplicate therapy and displayed safety warnings.",
    "Use Tests ordered to search the laboratory catalog or enter an appropriate free-text test.",
], numbered=True)

new_content_page("Doctor workflow", "Record imaging and control operational orders", "The clinical recommendation fields remain available even when the hospital does not operate an internal laboratory or imaging worklist.")
add_screenshot("Doctor's Desk Patient Form - 02.png", "Imaging recommendations, disabled operational orders and visit signature", max_height=3.85)
add_bullets([
    ("Imaging ordered. ", "Search the catalog by procedure, modality or body part. The catalog includes 271 procedures across 12 modalities in this edition."),
    ("Tests ordered. ", "Search the laboratory catalog by name, profile, category or abbreviation. Free-text entry remains available."),
    ("Operational orders. ", "When disabled, the application records recommendations without creating departmental or HL7 work."),
    ("Signature. ", "Save a draft while documentation is incomplete. Confirm the signature and complete the visit only after reviewing the record."),
])
add_lead("Corrections after signing", "Use the hospital's approved amendment process so the original signed record remains traceable.")

# Records and diagnostics
new_content_page("Patient records", "Find the correct patient record", "Patient Records provides hospital-scoped search and access to permitted longitudinal information.")
add_screenshot("Patient Record - 01.png", "Patient Records search results", max_height=4.15)
add_bullets([
    "Search by MRN, patient name or phone.",
    "Verify more than one identifier before opening the record.",
    "If two records may belong to the same person, stop and use the approved duplicate-review process.",
], numbered=True)

new_content_page("Patient records", "Review the longitudinal clinical summary", "The selected record brings together demographics, safety information, encounters, medication history and other permitted clinical information.")
add_screenshot("Patient Record - 02.png", "Expanded patient record and clinical summary", max_height=4.15)
add_bullets([
    "Review allergies and active safety information before clinical decisions.",
    "Use the encounter history to understand previous diagnoses, medicines, tests and imaging recommendations.",
    "Edit demographics only after confirming the information with the patient and following hospital policy.",
])
add_lead("Minimum necessary access", "Open a patient record only when it is needed for assigned work. Patient access may be recorded in audit history.")

new_content_page("Diagnostics", "Process operational diagnostic orders", "The Diagnostics workspace is used only for orders created through the enabled operational workflow.")
add_screenshot("Diagnostics.png", "Diagnostics worklist with report processing and completed orders", max_height=4.15)
add_bullets([
    "Filter the worklist by type or status when needed.",
    "Begin collection or processing only when the sample or patient reaches the department.",
    "Complete the work in the hospital's approved laboratory or radiology process.",
    "Attach the final approved report and send it to the ordering doctor.",
    "The ordering clinician opens and acknowledges the completed report after review.",
], numbered=True)
add_lead("Status meaning", "Completed means a report was uploaded. Reviewed means an authorized clinician acknowledged it. Neither status replaces clinical review of the report.")

new_content_page("Documents", "Store approved patient documents", "Documents keeps supported patient files inside the selected hospital workspace with controlled access.")
add_screenshot("Documents.png", "Patient document search and upload workspace", max_height=4.15)
add_bullets([
    "Search for the patient and verify the correct record.",
    "Choose an approved PDF, JPEG or PNG file.",
    "Link a final report to its diagnostic order when applicable.",
    "Confirm the patient, document type and active hospital before upload.",
], numbered=True)
add_lead("Upload safety", "Do not upload unrelated files, passwords, multi-patient exports or documents that have not been approved under hospital policy.")

# Billing and audit
new_content_page("Billing", "Review invoices and outstanding balances", "Billing lists invoices for the active hospital and provides access to the full invoice workflow.")
add_screenshot("Billing - 01.png", "Billing list with invoice status and balance", max_height=4.2)
add_bullets([
    "Open the correct patient invoice after confirming the visit and invoice date.",
    "Review status, total, amount paid and remaining balance.",
    "Investigate unexpected zero prices before collecting payment.",
], numbered=True)

new_content_page("Billing", "Finalize invoice items and payments", "Invoice prices are populated from the active clinic's configured service catalog and must be reviewed before payment is recorded.")
add_screenshot("Billing - 02.png", "Invoice detail with line items, discounts, tax and totals", max_height=3.35)
add_bullets([
    "Review every line item, quantity, unit price, discount, tax and total.",
    "Save changes only after the invoice matches the approved services.",
    "Record the payment method and reference used by the hospital.",
    "For refunds or voids, enter a clear reason so the original transaction remains traceable.",
])
add_screenshot("Billing - 03.png", "Payment and refund history", max_height=1.3)

new_content_page("Audit", "Review important activity", "Audit history provides an append-only view of important security and workflow events for the active hospital.")
add_screenshot("Audit.png", "Audit history with time, action, resource, actor and reason", max_height=4.1)
add_bullets([
    "Use audit history only for authorized operational, privacy, security and support reviews.",
    "Search by event details when investigating a known action or time period.",
    "Do not place passwords, tokens or unnecessary clinical information in an audit reason.",
])
add_lead("Interpretation", "An audit event shows that an action occurred. It does not by itself determine whether that action was clinically appropriate.")

# Portal
new_content_page("Patient portal", "Use the separate patient sign in", "Patients use the patient portal rather than the staff sign-in page.")
add_screenshot("Patient Portal Login.png", "Patient portal sign-in page", max_height=4.3)
add_bullets([
    "The patient enters the registered email address and portal password.",
    "Staff should direct patients to the approved password-reset or support process when access is lost.",
    "A portal account must remain linked only to the correct patient record.",
])

new_content_page("Patient portal", "Activate patient access", "Portal activation should follow the hospital's identity-verification process before an account is created.")
add_screenshot("Patient Portal Register.png", "Patient portal activation form", max_height=4.3)
add_bullets([
    "Confirm the patient's MRN, registered email and date of birth.",
    "Ask the patient to create a private password that follows the displayed requirements.",
    "Do not activate access using an unverified email address or another person's device without consent.",
], numbered=True)

new_content_page("Patient portal", "Review the patient dashboard", "After sign-in, the patient sees only the information connected to the activated patient account.")
add_screenshot("Patient Portal - 01.png", "Patient dashboard with details and clinical summary", max_height=4.1)
add_bullets([
    "The dashboard summarizes demographics, clinical information and recent activity available to the patient.",
    "Patients should contact the hospital through the approved channel if information appears incorrect.",
    "Portal access does not replace urgent or emergency clinical communication.",
])

new_content_page("Patient portal", "View records and billing history", "The patient portal presents available records and billing information without exposing the staff workspace.")
add_screenshot("Patient Portal - 02.png", "Patient records and billing history", max_height=3.15)
add_bullets([
    "Open the relevant section to review available record or invoice details.",
    "Use the hospital's approved process for corrections, copies or billing questions.",
    "Sign out after using a shared or public device.",
])

# Operational pages
new_content_page("Go live", "New hospital setup checklist", "Complete and approve each item before entering real patient information.")
add_bullets([
    "Hospital account created and selected as the active hospital",
    "Clinic locations, codes, addresses, timezones, departments and appointment types entered",
    "Doctor accounts linked to practitioner profiles and weekly schedules verified",
    "Breaks, blocked periods and clinic holidays tested",
    "Operational diagnostic and imaging capabilities enabled only when integrations are ready",
    "Service catalog, clinic prices and tax rules approved by the hospital's finance owner",
    "Staff accounts assigned the minimum required roles",
    "Email delivery, document storage, encryption, malware scanning, HTTPS and audit review tested",
    "Encrypted backups completed and a restore test documented",
    "A complete fictional-patient workflow successfully tested",
], numbered=False)
add_lead("Approval", "Clinical, operational, privacy, legal, security and finance owners should approve the configured workflow before real use.")

new_content_page("Daily operations", "Daily workflow checklists")
doc.add_paragraph("Walk-in visit", style="Heading 2")
add_bullets([
    "Front Desk searches for or registers the patient.",
    "Front Desk starts the visit and selects the doctor.",
    "Doctor opens the patient from Waiting Room.",
    "Doctor documents the consultation, recommendations and any enabled operational orders.",
    "Lab or radiology processes operational work and uploads the final report when applicable.",
    "Ordering clinician reviews the report.",
    "Billing reviews the invoice and records payment.",
], numbered=True)
doc.add_paragraph("Online appointment", style="Heading 2")
add_bullets([
    "Patient submits a public appointment request.",
    "Authorized staff or the doctor reviews and confirms it.",
    "Front Desk verifies and checks the patient in on arrival.",
    "The clinical, diagnostic and billing workflow continues as for a walk-in visit.",
], numbered=True)

new_content_page("Privacy and security", "Protect patient and hospital information", "Every user is responsible for handling patient information only for approved hospital work.")
add_bullets([
    "Use an individual account and lock the device when stepping away.",
    "Confirm the active hospital and patient before viewing, entering, uploading, printing or billing information.",
    "Access only the minimum information needed for assigned work.",
    "Use fictional identities in demonstrations, training materials and sales presentations.",
    "Remove patient details, staff emails, internal identifiers, tokens and private URLs from support screenshots.",
    "Follow the hospital's approved retention, backup, access-review and incident-response procedures.",
    "Apply the legal and regulatory requirements relevant to the hospital and deployment location.",
])
add_lead("Important", "This guide explains product workflows. It is not a certification of regulatory compliance, clinical safety or production readiness.")

new_content_page("Support", "Troubleshooting and safe escalation")
add_table(
    ["Problem", "First checks"],
    [
        ["Wrong hospital is shown", "Open Active Hospital if authorized, or confirm assignment to the correct hospital."],
        ["A menu is missing", "Check the signed-in role; functions outside that role are hidden."],
        ["Patient is not found", "Confirm the hospital and search another approved identifier before creating a record."],
        ["Doctor is missing", "Confirm an active doctor account is linked to a practitioner profile and clinic."],
        ["No appointment slots", "Check schedule, breaks, blocked time, holidays and existing appointments."],
        ["Invoice price is zero", "Verify the service name and active clinic price in Settings."],
        ["Report cannot be uploaded", "Confirm the order status and use an approved file type within the size limit."],
        ["Page remains loading", "Refresh once, then contact support with the page name and approximate time."],
    ],
    [2.0, 5.05],
)
doc.add_paragraph("Information safe to send support", style="Heading 2")
add_bullets([
    "Page name and approximate time of the issue",
    "Your role and the hospital name shown in the header",
    "The non-sensitive error message and steps that led to it",
    "A screenshot only after patient details, credentials, tokens, emails, identifiers and private URLs are removed",
])
add_lead("Do not send", "Passwords, authentication codes, database credentials, complete patient exports or unapproved clinical documents.")

doc.save(OUT)
print(OUT)
print(f"screens embedded: {len(doc.inline_shapes)}")
print(f"paragraphs: {len(doc.paragraphs)} tables: {len(doc.tables)}")
